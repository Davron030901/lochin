"""Multimodal ingestion pipeline: Load -> Caption images -> Chunk -> Embed -> Store.

Supported inputs, each routed to the right loader:
    * PDF (.pdf)                 -> PyMuPDF text + embedded-image captioning
    * plain text (.txt/.md)      -> load text, chunk directly
    * Word (.docx)               -> text (paragraphs + tables) + embedded images
    * standalone image (.png/.jpg/.jpeg/.webp) -> caption directly, index as
      an "[Image] ..." document (same treatment as PDF-embedded images)

Images are captioned by a vision LLM and added as their own ``[Image] ...``
documents so diagrams/screenshots become retrievable alongside text.

Idempotency: each chunk gets a deterministic point-id derived from a content
hash (+ source + index), so re-ingesting the same file upserts in place instead
of creating duplicates.
"""

from __future__ import annotations

import base64
import hashlib
import os
import uuid
import zipfile
from dataclasses import dataclass, field
from typing import Iterable, List, Optional, Tuple

from langchain_core.documents import Document
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_text_splitters import RecursiveCharacterTextSplitter

from .config import get_settings
from .llms import get_vision_model
from .prompts import CAPTION_SYSTEM, CAPTION_USER
from .vectorstore import get_vector_store

# Qdrant's namespace for building deterministic UUID5 point-ids.
_ID_NAMESPACE = uuid.UUID("12345678-1234-5678-1234-567812345678")

# Extension -> logical file type reported to the frontend.
_TEXT_EXTS = {".txt", ".md", ".markdown"}
_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp"}
_PDF_EXTS = {".pdf"}
_DOCX_EXTS = {".docx"}
SUPPORTED_EXTS = _TEXT_EXTS | _IMAGE_EXTS | _PDF_EXTS | _DOCX_EXTS

_IMAGE_MIME = {
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".webp": "image/webp",
    ".gif": "image/gif",
}


class UnsupportedFormatError(ValueError):
    """Raised when a file extension is not supported."""


@dataclass
class FileResult:
    """Structured per-file ingestion outcome (mirrors the API response)."""

    filename: str
    file_type: str  # "pdf" | "docx" | "text" | "image"
    chunks: int = 0
    images_captioned: int = 0
    pages: Optional[int] = None
    status: str = "indexed"  # "indexed" | "error"
    detail: str = ""


@dataclass
class IngestResult:
    """Aggregate summary (used by directory ingestion / eval)."""

    ingested_files: List[str] = field(default_factory=list)
    text_chunks: int = 0
    image_captions: int = 0
    points_upserted: int = 0
    files: List[FileResult] = field(default_factory=list)


def file_type_for(ext: str) -> str:
    ext = ext.lower()
    if ext in _PDF_EXTS:
        return "pdf"
    if ext in _DOCX_EXTS:
        return "docx"
    if ext in _IMAGE_EXTS:
        return "image"
    if ext in _TEXT_EXTS:
        return "text"
    raise UnsupportedFormatError(f"Unsupported file type: {ext}")


def is_supported(filename: str) -> bool:
    return os.path.splitext(filename)[1].lower() in SUPPORTED_EXTS


def _deterministic_id(source: str, index: int, content: str) -> str:
    """Stable UUID5 id from source+index+content hash -> idempotent upserts."""
    digest = hashlib.sha256(content.encode("utf-8")).hexdigest()
    return str(uuid.uuid5(_ID_NAMESPACE, f"{source}::{index}::{digest}"))


def caption_image(image_bytes: bytes, mime: str = "image/png") -> str:
    """Caption a single image with the vision model.

    Returns a short factual description suitable for retrieval. On failure it
    returns an empty string so ingestion degrades gracefully rather than
    aborting the whole document.
    """
    b64 = base64.b64encode(image_bytes).decode("utf-8")
    try:
        model = get_vision_model()
        message = HumanMessage(
            content=[
                {"type": "text", "text": CAPTION_USER},
                {
                    "type": "image_url",
                    "image_url": {"url": f"data:{mime};base64,{b64}"},
                },
            ]
        )
        resp = model.invoke([SystemMessage(content=CAPTION_SYSTEM), message])
        return (resp.content or "").strip()
    except Exception:  # pragma: no cover - provider/runtime dependent
        return ""


def _image_doc(source: str, caption: str, **extra) -> Document:
    meta = {"source": source, "modality": "image"}
    meta.update(extra)
    return Document(page_content=f"[Image] {caption}", metadata=meta)


# ---------------------------------------------------------------------------
# Per-format loaders. Each returns (text_docs, image_docs, pages).
# ---------------------------------------------------------------------------
def _load_pdf(path: str) -> Tuple[List[Document], List[Document], Optional[int]]:
    import fitz  # PyMuPDF (imported lazily so tests don't require it)

    text_docs: List[Document] = []
    image_docs: List[Document] = []
    source = os.path.basename(path)

    with fitz.open(path) as doc:
        pages = len(doc)
        for page_index in range(pages):
            page = doc[page_index]
            text = page.get_text("text").strip()
            if text:
                text_docs.append(
                    Document(
                        page_content=text,
                        metadata={"source": source, "page": page_index + 1, "modality": "text"},
                    )
                )
            for img_index, img in enumerate(page.get_images(full=True)):
                xref = img[0]
                try:
                    base = doc.extract_image(xref)
                except Exception:  # pragma: no cover
                    continue
                image_bytes = base.get("image")
                if not image_bytes:
                    continue
                mime = f"image/{base.get('ext', 'png')}"
                caption = caption_image(image_bytes, mime=mime)
                if not caption:
                    continue
                image_docs.append(
                    _image_doc(source, caption, page=page_index + 1, image_index=img_index)
                )
    return text_docs, image_docs, pages


def _load_text(path: str) -> Tuple[List[Document], List[Document], Optional[int]]:
    with open(path, "r", encoding="utf-8", errors="ignore") as fh:
        content = fh.read()
    return (
        [Document(page_content=content, metadata={"source": os.path.basename(path), "modality": "text"})],
        [],
        None,
    )


def _load_docx(path: str) -> Tuple[List[Document], List[Document], Optional[int]]:
    """Extract paragraph + table text and any embedded images from a .docx."""
    import docx  # python-docx (lazy import)

    source = os.path.basename(path)
    document = docx.Document(path)

    parts: List[str] = [p.text for p in document.paragraphs if p.text and p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append("\t".join(cells))
    text = "\n".join(parts).strip()

    text_docs: List[Document] = []
    if text:
        text_docs.append(
            Document(page_content=text, metadata={"source": source, "modality": "text"})
        )

    # Embedded images live under word/media/* inside the .docx zip container.
    image_docs: List[Document] = []
    try:
        with zipfile.ZipFile(path) as zf:
            media = [n for n in zf.namelist() if n.startswith("word/media/")]
            for idx, name in enumerate(sorted(media)):
                ext = os.path.splitext(name)[1].lower()
                if ext not in _IMAGE_MIME:
                    continue
                data = zf.read(name)
                caption = caption_image(data, mime=_IMAGE_MIME[ext])
                if caption:
                    image_docs.append(_image_doc(source, caption, image_index=idx))
    except zipfile.BadZipFile as exc:  # corrupt docx
        raise ValueError("The .docx file appears to be corrupt or unreadable.") from exc

    return text_docs, image_docs, None


def _load_image(path: str) -> Tuple[List[Document], List[Document], Optional[int]]:
    """Caption a standalone image and index the caption as an [Image] document."""
    source = os.path.basename(path)
    ext = os.path.splitext(path)[1].lower()
    with open(path, "rb") as fh:
        data = fh.read()
    caption = caption_image(data, mime=_IMAGE_MIME.get(ext, "image/png"))
    if not caption:
        raise ValueError("The vision model could not caption this image.")
    return [], [_image_doc(source, caption)], None


def _load_file(path: str) -> Tuple[List[Document], List[Document], Optional[int], str]:
    """Dispatch a single file to the right loader; return (text, images, pages, file_type)."""
    ext = os.path.splitext(path)[1].lower()
    ftype = file_type_for(ext)  # raises UnsupportedFormatError for unknown types
    if ftype == "pdf":
        t, i, p = _load_pdf(path)
    elif ftype == "docx":
        t, i, p = _load_docx(path)
    elif ftype == "image":
        t, i, p = _load_image(path)
    else:  # text
        t, i, p = _load_text(path)
    return t, i, p, ftype


def _chunk_and_upsert(docs: List[Document]) -> int:
    """Chunk documents and upsert into the collection; return chunk count."""
    if not docs:
        return 0
    settings = get_settings()
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size, chunk_overlap=settings.chunk_overlap
    )
    chunks = splitter.split_documents(docs)
    ids = [
        _deterministic_id(str(c.metadata.get("source", "?")), i, c.page_content)
        for i, c in enumerate(chunks)
    ]
    get_vector_store().add_documents(chunks, ids=ids)
    return len(chunks)


# ---------------------------------------------------------------------------
# Public API.
# ---------------------------------------------------------------------------
def ingest_file(path: str) -> FileResult:
    """Ingest exactly one file and return a structured per-file result."""
    filename = os.path.basename(path)
    ext = os.path.splitext(path)[1].lower()
    if ext not in SUPPORTED_EXTS:
        raise UnsupportedFormatError(f"Unsupported format: {ext or '(none)'}")

    text_docs, image_docs, pages, ftype = _load_file(path)
    chunk_count = _chunk_and_upsert(text_docs + image_docs)
    return FileResult(
        filename=filename,
        file_type=ftype,
        chunks=chunk_count,
        images_captioned=len(image_docs),
        pages=pages,
        status="indexed",
    )


def ingest_paths(path: str) -> IngestResult:
    """Ingest a single file or every supported file in a directory."""
    result = IngestResult()
    paths: Iterable[str]
    if os.path.isdir(path):
        collected = []
        for root, _dirs, files in os.walk(path):
            for name in sorted(files):
                if is_supported(name):
                    collected.append(os.path.join(root, name))
        paths = collected
    else:
        paths = [path]

    for file_path in paths:
        try:
            fr = ingest_file(file_path)
        except UnsupportedFormatError:
            continue
        result.files.append(fr)
        result.ingested_files.append(fr.filename)
        result.text_chunks += fr.chunks
        result.image_captions += fr.images_captioned
        result.points_upserted += fr.chunks
    return result


def ingest_bytes(filename: str, data: bytes) -> FileResult:
    """Ingest raw uploaded bytes (single file) -> structured per-file result."""
    ext = os.path.splitext(filename)[1].lower()
    if ext not in SUPPORTED_EXTS:
        raise UnsupportedFormatError(f"Unsupported format: {ext or '(none)'}")

    tmp_dir = os.path.join(get_settings().qdrant_path, "..", "uploads_tmp")
    os.makedirs(tmp_dir, exist_ok=True)
    tmp_path = os.path.join(tmp_dir, filename)
    with open(tmp_path, "wb") as fh:
        fh.write(data)
    try:
        return ingest_file(tmp_path)
    finally:
        try:
            os.remove(tmp_path)
        except OSError:  # pragma: no cover
            pass
