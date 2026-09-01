"""Format-routing tests for the ingestion pipeline.

The vision captioner and the vector store are mocked, so these run with no API
keys and no real Qdrant. They verify that each format is routed to the right
loader and produces the correct structured ``FileResult`` (the shape the API and
frontend depend on).
"""

from __future__ import annotations

import base64
import os

import pytest

from app import ingestion
from app.ingestion import UnsupportedFormatError, ingest_file

# 1x1 transparent PNG.
_PNG_1x1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+M9QDwADhgGAWjR9awAAAABJRU5ErkJggg=="
)


class _FakeStore:
    """Captures documents that would have been upserted."""

    def __init__(self):
        self.added = []

    def add_documents(self, chunks, ids=None):
        self.added.extend(chunks)


@pytest.fixture
def fake_store(monkeypatch):
    store = _FakeStore()
    monkeypatch.setattr(ingestion, "get_vector_store", lambda: store)
    # Deterministic caption without any vision/LLM call.
    monkeypatch.setattr(ingestion, "caption_image", lambda data, mime="image/png": "a red square")
    return store


def test_text_file_routing(tmp_path, fake_store):
    p = tmp_path / "notes.txt"
    p.write_text("Paris is the capital of France. " * 20, encoding="utf-8")

    result = ingest_file(str(p))

    assert result.file_type == "text"
    assert result.chunks >= 1
    assert result.images_captioned == 0
    assert result.pages is None
    assert result.status == "indexed"
    assert any("Paris" in d.page_content for d in fake_store.added)


def test_markdown_routes_as_text(tmp_path, fake_store):
    p = tmp_path / "readme.md"
    p.write_text("# Title\n\nSome markdown body.", encoding="utf-8")
    result = ingest_file(str(p))
    assert result.file_type == "text"
    assert result.chunks >= 1


def test_standalone_image_routing(tmp_path, fake_store):
    p = tmp_path / "diagram.png"
    p.write_bytes(_PNG_1x1)

    result = ingest_file(str(p))

    assert result.file_type == "image"
    assert result.images_captioned == 1
    assert result.chunks == 1
    added = fake_store.added
    assert len(added) == 1
    assert added[0].page_content.startswith("[Image] a red square")
    assert added[0].metadata["modality"] == "image"


def test_jpeg_extension_routes_as_image(tmp_path, fake_store):
    p = tmp_path / "photo.jpeg"
    p.write_bytes(_PNG_1x1)  # bytes content irrelevant; captioner is mocked
    result = ingest_file(str(p))
    assert result.file_type == "image"
    assert result.images_captioned == 1


def test_docx_routing_text_and_images(tmp_path, fake_store):
    import docx

    d = docx.Document()
    d.add_paragraph("The proposed method improves retrieval accuracy.")
    d.add_paragraph("It was evaluated on a benchmark dataset.")
    path = tmp_path / "paper.docx"
    d.save(str(path))

    result = ingest_file(str(path))

    assert result.file_type == "docx"
    assert result.chunks >= 1
    assert any("proposed method" in doc.page_content for doc in fake_store.added)


def test_unsupported_format_raises(tmp_path, fake_store):
    p = tmp_path / "archive.zip"
    p.write_bytes(b"PK\x03\x04not-a-real-zip")
    with pytest.raises(UnsupportedFormatError):
        ingest_file(str(p))


def test_uncaptionable_image_raises_readable_error(tmp_path, monkeypatch):
    monkeypatch.setattr(ingestion, "get_vector_store", lambda: _FakeStore())
    monkeypatch.setattr(ingestion, "caption_image", lambda data, mime="image/png": "")
    p = tmp_path / "blank.png"
    p.write_bytes(_PNG_1x1)
    with pytest.raises(ValueError) as exc:
        ingest_file(str(p))
    assert "caption" in str(exc.value).lower()
